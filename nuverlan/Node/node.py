
from nuverlan.llm.model import model
from docx import Document
from nuverlan.Utils.util import generate_comprehensive_spec
from langchain_core.tools import tool
from langgraph.types import Command, interrupt
from langchain_core.messages import HumanMessage
import websockets
from nuverlan.websocket.websocket import WebSocketHandler
import json
import asyncio
class node:
    def __init__(self):
        self.model = model()
        self.websocket_handler = WebSocketHandler()
        # https://github.com/gotohuman/examples-langgraph-py/tree/main/sales-lead-agent
  
    def human_assistance(self,query: str) -> str:
        '''Request assistance from a human.'''
        print('before')
        location = interrupt('Please provide your location:')
        tool_message ='Approved'
        return 'Approved'
    
    def generate_user_stories(self,state):
        print('***********Generate User Story***********')
        prompt = '''
You are an AI assistant specializing in front-end development with React.js.  
Create a complete and fully functional swiggy application using React.js based on the following requirements.  
Return the response in a **structured JSON format** to facilitate creating folders, configurations, and files for executing the application.  

### Requirements:
1. **Core Features:** 
   - Product listing with search, filter, and sorting options.
   - Product detail pages with descriptions, images, pricing, and availability.
   - Add to cart, view cart, update quantity, and remove items.
   - Checkout process with form validation and order confirmation.

2. **User Management:** 
   - User authentication: registration, login, and logout.
   - User profile management: view and edit personal information, order history.

3. **Integration:** 
   - Import existing product data from a provided JSON file or database.
   - Integrate user authentication using existing credentials or an external API.
   - Payment gateway integration for secure transactions (e.g., Stripe).

4. **State Management:** 
   - Efficient state management using React Context API, Redux, or Zustand.

5. **UI/UX Design:** 
   - Modern, responsive, and accessible design with styled-components, Tailwind CSS, or Material-UI.
   - Clear navigation and user-friendly interface.

6. **Folder Structure:** 
   - Follow best practices for organizing components, hooks, contexts, and assets.
   - Include configuration files and environment setup for production deployment.
   - Import existing configuration files, if available.

7. **Error Handling and Security:** 
   - Form validation, error boundaries, and secure API requests.

8. **Documentation:** 
   - Provide a README file with setup instructions, scripts, and deployment guidance.
   - Explain how to integrate and use the existing data files.

Only refer the below JSON response as example  and give dynamic filename and folder based on the application
### JSON Response Format Example:
```json
{
  'project': {
    'name': 'RealEstateApp',
    'folders': [
      {
        'name': 'public',
        'files': [
          {
            'name': 'index.html',
            'content': '<!DOCTYPE html><html lang=\'en\'><head><meta charset=\'UTF-8\'><meta name=\'viewport\' content=\'width=device-width, initial-scale=1.0\'><title>Real Estate App</title></head><body><div id=\'root\'></div><script src=\'index.js\'></script></body></html>'
          }
        ]
      },
      {
        'name': 'src',
        'folders': [
          {
            'name': 'components',
            'files': [
              {
                'name': '<filename>.js',
                'content': '// React component for product listing'
              },
              {
                'name': '<filename>.js',
                'content': '// React component for product detail'
              },
              {
                'name': '<filename>.js',
                'content': '// React component for cart'
              },
              {
                'name': '<filename>.js',
                'content': '// React component for checkout'
              },
              {
                'name': '<filename>.js',
                'content': '// React component for user profile'
              }
            ]
          },
          {
            'name': 'contexts',
            'files': [
              {
                'name': '<filename>.js',
                'content': '// React context for product data'
              },
              {
                'name': '<filename>.js',
                'content': '// React context for user data'
              }
            ]
          },
          {
            'name': 'hooks',
            'files': [
              {
                'name': '<filename>.js',
                'content': '// Custom hook for product data'
              },
              {
                'name': '<filename>.js',
                'content': '// Custom hook for user data'
              }
            ]
          },
          {
            'name': 'utils',
            'files': [
              {
                'name': '<filename>.js',
                'content': '// API utility functions'
              },
              {
                'name': '<filename>.js',
                'content': '// Storage utility functions'
              }
            ]
          },
          {
            'name': 'assets',
            'files': [
              {
                'name': '<filename>.json',
                'content': '// Existing product data'
              }
            ]
          },
          {
            'name': 'styles',
            'files': [
              {
                'name': '<filename>.css',
                'content': '// Global stylesheet'
              },
              {
                'name': '<filename>.css',
                'content': '// Component-specific stylesheet'
              }
            ]
          }
        ]
      },
      {
        'name': 'config',
        'files': [
          {
            'name': 'env.js',
            'content': '// Environment configuration'
          },
          {
            'name': 'stripe.js',
            'content': '// Stripe payment gateway configuration'
          }
        ]
      }
    ]
  },
  'scripts': [
    {
      'name': 'start',
      'command': 'react-scripts start'
    },
    {
      'name': 'build',
      'command': 'react-scripts build'
    },
    {
      'name': 'deploy',
      'command': 'react-scripts deploy'
    }
  ],
  'dependencies': {
    'react': '^17.0.2',
    'react-dom': '^17.0.2',
    'react-scripts': '4.0.3',
    'styled-components': '^5.3.3',
    'tailwindcss': '^3.0.2',
    'stripe-js': '^1.24.0'
  },
  'README': {
    'content': '# Real Estate App\n\n## Setup\n\n1. Clone the repository\n2. Run `npm install` to install dependencies\n3. Run `npm start` to start the development server\n\n## Deployment\n\n1. Run `npm run build` to build the application\n2. Run `npm run deploy` to deploy the application\n\n## Integration\n\n1. Import existing product data from `product-data.json`\n2. Integrate user authentication using existing credentials or an external API\n3. Integrate Stripe payment gateway for secure transactions'
  }
}
Return the 1 user stories as an array of objects, where each object represents a separate user story.
'''


        #prompt = f'Generate a list of 3 user stories for an {state['user_requirement']}, with each user story formatted as: 'As a [type of user], I want [goal] so that [reason].' Please return the response as a list (array format).'
        data = self.model.stream_llm_response(prompt)
        print('*************************************')
        # print(data.split('**Folder Structure:**'))
        # print('+++++++++++++++++++++++++++')
        user_story_processed_data = data.split("\n```json")[2]
        lines = user_story_processed_data.splitlines()

        if lines[-1].startswith("```"):
            user_story_processed_data = "\n".join(lines[:-1])  # Exclude the first and last lines
        else:
            user_story_processed_data = user_story_processed_data  # No changes if backticks are not present
        print(user_story_processed_data)
        print("?????????????")
        #websockets.send_text(data['user_story'])
        # p = ""
        # loop = asyncio.get_event_loop()
        # loop.run_in_executor(None, self.websocket_handler.send_message('1', f"user_story:{user_story_processed_data}"), p)
        
        # return {'user_story':[self.model.process().invoke(state['user_requirement'] + 'Create the list of user stories with acceptence critera')]}
        return {'user_story':user_story_processed_data}

    def product_owner(self,state):
        print('***********Product Owner***********')
        review = interrupt('Review the user story:')
        print(review)
        return {'product_owner': review}

    def product_owner_status(self,state):
       print('***********Product Owner Status***********')
       
       return state['product_owner']
    
    def create_design_document(self,state):
        print('***********create Design Document***********')
        # comprehensive_spec = generate_comprehensive_spec(state['user_story'])
        
        # doc = Document()
        # doc.add_heading('Design Documentation', 0)

        # # Add the generated content to the Word document
        # doc.add_paragraph(comprehensive_spec)

        # # Save the DOCX document
        # doc.save('comprehensive_specification.docx')
        return {'design_document':'Created'}
       


    def revise_user_stories(self,state):
        print('***********Revise User Story***********')



    def design_review(self,state):
        print('***********Review Design***********')

    def validate_design_review(self,state):
        print('***********VALIDATE DESIGN REVIEW***********')
        status = input('Validate the document?')
        return status

    def generate_code(self,state):
        print('***********Generate Code***********')
        return state  
        #return "Approved" if state["code_review_files"].length==0 else "FeedBack"  
    
    def assistant(self,state):
        return {'messages': [self.model.process().invoke(state['messages'])]}
    
    def code_review(self,state):
        review_status = []
        for path in state["generated_files"]:
            with open(path, "r") as file:
              content = file.read()

            status = self.code_review_details(content,path)
            json_data = json.loads(status)
            if json_data['ReviewRequired']:
              json_data["File"] = path
              review_status.append(json_data)
        print("************CODE REVIEW*************")
        print(review_status)
        return {"code_review_files": [*review_status]}
        

    def code_review_details(self,code,path):
        prompt="""
                  Review the provided code carefully. Identify issues or improvements needed and respond **ONLY** with a JSON object in the following format:
                    {
                    "File": "Name of the file being reviewed (e.g., models.py, views.py, auth.py)",
                    "IssuesFound": [
                    {
                    "IssueType": "Type of issue (e.g., Security, Logic Error, Code Style, Performance)",
                    "Description": "Detailed description of the issue or improvement needed",
                    "SuggestedFix": "Suggested solution or improvement"
                    }
                    ],
                    "ReviewRequired": true or false
                    }
                  Review the following code:

                """
        print("***********CODE REVIEW ******************")
        print(path)
        data = self.model.stream_llm_response(prompt+code)
        return data
    
    def code_review_status(self,state):
        return "Approved" if state["code_review_files"].length>0 else "FeedBack"
    
    def fix_code(self,state):
        for path in state["code_review_files"]:
            with open(path['File'], "r") as file:
              content = file.read()

            status = self.code_review_details(content,path)
            json_data = json.loads(status)
            if json_data['ReviewRequired']:
              json_data["File"] = path
        state["code_review_files"] = []
        return {"code_review_files": []}